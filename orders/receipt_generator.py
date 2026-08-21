from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from decimal import Decimal
from datetime import datetime
import json
from io import BytesIO
from .models import Order
from frames.services import PriceCalculator, OrderExtrasCalculator


def add_table_border(table):
    """Добавляет границы к таблице"""
    tbl = table._tbl
    tblBorders = OxmlElement('w:tblBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    tbl.tblPr.append(tblBorders)


def format_number(value):
    """Форматирует число с правильным количеством знаков после запятой"""
    if value is None:
        return "—"
    try:
        if isinstance(value, Decimal):
            val = float(value)
        else:
            val = float(value)
        
        # Форматируем с 2 знаками после запятой, но убираем лишние нули
        formatted = f"{val:.2f}"
        if '.' in formatted:
            formatted = formatted.rstrip('0').rstrip('.')
        return formatted
    except (ValueError, TypeError):
        return "—"


def _add_independent_passepartouts(calculation, order, frames):
    """Добавляет независимые паспарту (passepartout_1, ...) в расчёт квитанции.

    Такие паспарту хранятся отдельным списком в frames_data и в цену заказа
    входят, но раньше не выводились в квитанции. Возвращает добавленную сумму.
    """
    from frames.models import Passepartout
    pps = []
    if frames and isinstance(frames[0], dict):
        emb = frames[0].get('passepartouts')
        if isinstance(emb, list):
            pps = [p for p in emb if isinstance(p, dict) and p.get('passepartout_id')]
    added = Decimal('0')
    for i, pp in enumerate(pps):
        try:
            obj = Passepartout.objects.get(pk=pp['passepartout_id'])
        except Passepartout.DoesNotExist:
            continue
        length = Decimal(str(pp.get('passepartout_length'))) if pp.get('passepartout_length') else order.x1
        width = Decimal(str(pp.get('passepartout_width'))) if pp.get('passepartout_width') else order.x2
        area = PriceCalculator.calculate_glass_area(length, width)
        price = area * obj.price
        calculation['components'][f'passepartout_{i + 1}'] = {
            'name': f'{obj.name} (Паспарту {i + 1})',
            'length': float(length),
            'width': float(width),
            'area': float(area),
            'unit_price': float(obj.price),
            'total_price': float(price),
        }
        calculation['total_price'] += Decimal(str(price))
        added += Decimal(str(price))
    return added


def _order_backing_ids(order, frames):
    """Список id подкладок из frames_data (несколько), иначе одна из заказа."""
    ids = []
    if frames and isinstance(frames[0], dict):
        emb = frames[0].get('backings')
        if isinstance(emb, list):
            ids = [(b.get('backing_id') if isinstance(b, dict) else b) for b in emb if b]
    if not ids and order.backing_id:
        ids = [order.backing_id]
    return ids


def _scale_calc_by_copies(calculation, q):
    """Умножает расход/стоимость материалов на количество копий (для детализации)."""
    q = int(q or 1) or 1
    if q <= 1:
        return
    for comp in calculation.get('components', {}).values():
        for k in ('total_price', 'area', 'quantity', 'consumption', 'length'):
            v = comp.get(k)
            if isinstance(v, (int, float)):
                comp[k] = v * q
    tp = calculation.get('total_price', 0)
    calculation['total_price'] = tp * q


def _scale_extras_by_copies(extras, q):
    """Умножает стоимость работ и сложность на количество копий."""
    q = int(q or 1) or 1
    if q <= 1:
        return
    works = extras.get('works', {})
    for w in works.get('items', []):
        w['total'] = w['total'] * q
    if 'total_rate' in works:
        works['total_rate'] = works['total_rate'] * q
    if extras.get('manual_complexity'):
        extras['manual_complexity'] = extras['manual_complexity'] * q


def add_horizontal_line(doc):
    """Добавляет горизонтальную пунктирную линию для отрыва"""
    para = doc.add_paragraph()
    para_format = para.paragraph_format
    para_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_format.space_before = Pt(6)
    para_format.space_after = Pt(3)
    
    # Создаем пунктирную линию
    run = para.add_run('─' * 50)
    run.font.size = Pt(9)
    
    # Добавляем текст "ОТРЫВНАЯ ЧАСТЬ"
    para2 = doc.add_paragraph()
    para2_format = para2.paragraph_format
    para2_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para2_format.space_before = Pt(0)
    para2_format.space_after = Pt(3)
    run2 = para2.add_run('ОТРЫВНАЯ ЧАСТЬ')
    run2.font.size = Pt(7)
    run2.font.bold = True
    run2.font.italic = True
    
    para3 = doc.add_paragraph()
    para3_format = para3.paragraph_format
    para3_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para3_format.space_before = Pt(0)
    para3_format.space_after = Pt(6)
    run3 = para3.add_run('─' * 50)
    run3.font.size = Pt(9)


def generate_receipt_word(order_id):
    """
    Генерирует Word документ квитанции для заказа
    Квитанция разделена на две части: отрывную для клиента и основную для организации
    
    Args:
        order_id: ID заказа
        
    Returns:
        BytesIO: Поток с содержимым Word документа
    """
    try:
        order = Order.objects.get(pk=order_id)
    except Order.DoesNotExist:
        raise ValueError(f"Заказ с ID {order_id} не найден")
    
    # Создаем документ
    doc = Document()
    
    # Настройка страницы (компактные поля)
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    
    # ========== ПЕРВАЯ ЧАСТЬ: ОТРЫВНАЯ ДЛЯ КЛИЕНТА ==========
    
    # Создаем компактную таблицу для отрывной части
    client_table = doc.add_table(rows=1, cols=2)
    client_table.style = 'Table Grid'
    
    # Левая колонка - основная информация
    left_cell = client_table.rows[0].cells[0]
    left_cell.vertical_alignment = 1  # Выравнивание по верху
    
    # Заголовок и номер заказа в одной строке
    date_str = order.created_at.strftime('%d.%m.%Y')
    title_para = left_cell.add_paragraph(f'Квитанция заказа № {order.pk} от {date_str}')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.size = Pt(11)
    title_run.font.bold = True
    
    # Сумма, аванс, долг компактно
    left_cell.add_paragraph()  # Небольшой отступ
    sum_para = left_cell.add_paragraph(f'Сумма: {format_number(order.total_price)} руб')
    sum_para.runs[0].font.size = Pt(10)
    sum_para.runs[0].font.bold = True
    
    advance_para = left_cell.add_paragraph(f'Аванс: {format_number(order.advance_payment)} руб')
    advance_para.runs[0].font.size = Pt(10)
    
    debt_para = left_cell.add_paragraph(f'Долг: {format_number(order.debt)} руб')
    debt_para.runs[0].font.size = Pt(10)
    
    # Правая колонка - информация об ИП (маленьким шрифтом)
    right_cell = client_table.rows[0].cells[1]
    right_cell.vertical_alignment = 1  # Выравнивание по верху
    
    ip_info = right_cell.add_paragraph()
    ip_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ip_run = ip_info.add_run('ИП Караковский С.М.\n')
    ip_run.font.size = Pt(6)
    ip_info.add_run('ОГРН ИП:304263203300215\n').font.size = Pt(6)
    ip_info.add_run('ИНН:263204326063\n').font.size = Pt(6)
    ip_info.add_run('Профессиональное оформление\n').font.size = Pt(6)
    ip_info.add_run('картин, постеров, фотографий,\n').font.size = Pt(6)
    ip_info.add_run('гобеленов и вышивок.\n').font.size = Pt(6)
    ip_info.add_run('г.Пятигорск, ул.Дзержинского, д.49А\n').font.size = Pt(6)
    ip_info.add_run('тел: 33-71-75, +7-962-41-30-631\n').font.size = Pt(6)
    ip_info.add_run('11:00 - 18:00. Выходной: понедельник.').font.size = Pt(6)
    
    # Убираем границы таблицы для более чистого вида
    tbl = client_table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    # Уменьшаем отступы в ячейках для компактности
    for row in client_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(1)
                para.paragraph_format.space_before = Pt(1)
    
    # Разделительная линия для отрыва
    add_horizontal_line(doc)
    
    # ========== ВТОРАЯ ЧАСТЬ: ДЛЯ ОРГАНИЗАЦИИ ==========
    
    # Компактная таблица с основной информацией
    info_table = doc.add_table(rows=1, cols=2)
    info_table.style = 'Table Grid'
    
    # Левая колонка - заголовок и номер заказа в одной строке
    left_info = info_table.rows[0].cells[0]
    left_info.vertical_alignment = 1
    
    title2_para = left_info.add_paragraph(f'Квитанция заказа № {order.pk} от {date_str}')
    title2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title2_run = title2_para.runs[0]
    title2_run.font.size = Pt(11)
    title2_run.font.bold = True
    
    # Срок изготовления если есть
    fulfillment_date_str = "—"
    if order.fulfillment_date:
        fulfillment_date_str = order.fulfillment_date.strftime('%d.%m.%Y')
    elif order.status in ['ready', 'issued']:
        fulfillment_date_str = order.updated_at.strftime('%d.%m.%Y')
    
    if fulfillment_date_str != "—":
        fulfillment_para = left_info.add_paragraph(f'Срок: {fulfillment_date_str}')
        fulfillment_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fulfillment_para.runs[0].font.size = Pt(9)
    
    # Правая колонка - данные клиента и финансы в одной строке
    right_info = info_table.rows[0].cells[1]
    right_info.vertical_alignment = 1
    
    # Формируем строку с информацией о клиенте, авансе и сумме
    customer_parts = []
    if order.customer_name:
        customer_parts.append(order.customer_name)
    if order.customer_phone:
        customer_parts.append(order.customer_phone)
    
    customer_info = 'Клиент: ' + ' '.join(customer_parts) if customer_parts else 'Клиент: —'
    customer_para = right_info.add_paragraph()
    
    # Добавляем части текста отдельными runs
    customer_para.add_run(f'{customer_info} | Аванс: {format_number(order.advance_payment)} руб | ').font.size = Pt(9)
    sum_run = customer_para.add_run(f'Сумма: {format_number(order.total_price)} руб')
    sum_run.font.size = Pt(9)
    sum_run.font.bold = True
    
    # Убираем границы и уменьшаем отступы
    tbl2 = info_table._tbl
    tblPr2 = tbl2.tblPr
    if tblPr2 is None:
        tblPr2 = OxmlElement('w:tblPr')
        tbl2.insert(0, tblPr2)
    
    tblBorders2 = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders2.append(border)
    tblPr2.append(tblBorders2)
    
    for row in info_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(1)
                para.paragraph_format.space_before = Pt(1)
    
    # Заголовок для таблицы с пунктами заказа (компактный)
    items_title = doc.add_paragraph('Пункты заказа:')
    items_title.runs[0].font.bold = True
    items_title.runs[0].font.size = Pt(10)
    items_title.paragraph_format.space_before = Pt(3)
    items_title.paragraph_format.space_after = Pt(2)
    
    # Получаем расчет стоимости
    frames = []
    if order.frames_data:
        try:
            frames = json.loads(order.frames_data)
        except (json.JSONDecodeError, TypeError):
            pass
    
    # Рассчитываем детализацию
    if frames and len(frames) > 0:
        calculation = {
            'components': {},
            'total_price': Decimal('0')
        }
        
        frame_sizes = []
        for idx, frame_data in enumerate(frames):
            if frame_data.get('baguette_id'):
                fx1 = Decimal(str(frame_data.get('x1', order.x1))) if frame_data.get('x1') else order.x1
                fx2 = Decimal(str(frame_data.get('x2', order.x2))) if frame_data.get('x2') else order.x2
                if not fx1 or not fx2 or fx1 <= 0 or fx2 <= 0:
                    fx1, fx2 = order.x1, order.x2
                frame_sizes.append((fx1, fx2))
                frame_calc = PriceCalculator.calculate_total_price(
                    x1=fx1,
                    x2=fx2,
                    baguette_id=frame_data.get('baguette_id'),
                    passepartout_id=frame_data.get('passepartout_id'),
                    passepartout_length=Decimal(str(frame_data.get('passepartout_length'))) if frame_data.get('passepartout_length') else None,
                    passepartout_width=Decimal(str(frame_data.get('passepartout_width'))) if frame_data.get('passepartout_width') else None,
                    work_id=frame_data.get('work_id'),
                )
                frame_num = idx + 1
                for key, value in frame_calc.get('components', {}).items():
                    if key in ['baguette', 'passepartout', 'work']:
                        component_key = f'{key}_frame{frame_num}'
                        calculation['components'][component_key] = {
                            **value,
                            'name': f"{value.get('name', key)} (Рама {frame_num})",
                        }
                        calculation['total_price'] += Decimal(str(value.get('total_price', 0)))
        if not frame_sizes:
            frame_sizes = [(order.x1, order.x2)]
        # Подкладка/подрамник — по раме меньшего размера (на картину, один раз)
        eff_x1, eff_x2 = min(frame_sizes, key=lambda s: s[0] * s[1])
        total_glass_area = sum(PriceCalculator.calculate_glass_area(fx1, fx2) for fx1, fx2 in frame_sizes)
        other_calc = PriceCalculator.calculate_total_price(
            x1=eff_x1,
            x2=eff_x2,
            glass_id=None,
            backing_id=order.backing.id if order.backing else None,
            backing_ids=_order_backing_ids(order, frames),
            hardware_id=order.hardware.id if order.hardware else None,
            hardware_quantity=order.hardware_quantity or 1,
            podramnik_id=order.podramnik.id if order.podramnik else None,
            package_id=order.package.id if order.package else None,
            package_quantity=order.package_quantity or 1,
            molding_id=order.molding.id if order.molding else None,
            molding_consumption=order.molding_consumption,
            trosik_id=order.trosik.id if order.trosik else None,
            trosik_length=order.trosik_length,
            podveski_id=order.podveski.id if order.podveski else None,
            podveski_quantity=order.podveski_quantity,
        )
        for key, value in other_calc.get('components', {}).items():
            if key not in ('glass', 'stretch'):
                calculation['components'][key] = value
                calculation['total_price'] += Decimal(str(value.get('total_price', 0)))
        if order.glass and total_glass_area > 0:
            glass = order.glass
            glass_calc = {
                'name': glass.name,
                'area': float(total_glass_area),
                'unit_price': float(glass.price_per_sqm),
                'total_price': float(total_glass_area * glass.price_per_sqm)
            }
            calculation['components']['glass'] = glass_calc
            calculation['total_price'] += Decimal(str(glass_calc['total_price']))
        # Независимые паспарту (отдельный список) — входят в цену, выводим их строками
        _add_independent_passepartouts(calculation, order, frames)
        # Натяжка — работа мастера (по периметру), выводится строкой «РАБОТА» из extras,
        # как материал по площади здесь не считается.
    else:
        calculation = PriceCalculator.calculate_total_price(
            x1=order.x1,
            x2=order.x2,
            baguette_id=order.baguette.id if order.baguette else None,
            glass_id=order.glass.id if order.glass else None,
            backing_id=order.backing.id if order.backing else None,
            backing_ids=_order_backing_ids(order, frames),
            hardware_id=order.hardware.id if order.hardware else None,
            hardware_quantity=order.hardware_quantity or 1,
            podramnik_id=order.podramnik.id if order.podramnik else None,
            package_id=order.package.id if order.package else None,
            package_quantity=order.package_quantity or 1,
            molding_id=order.molding.id if order.molding else None,
            molding_consumption=order.molding_consumption,
            trosik_id=order.trosik.id if order.trosik else None,
            trosik_length=order.trosik_length,
            podveski_id=order.podveski.id if order.podveski else None,
            podveski_quantity=order.podveski_quantity,
            stretch_id=order.stretch.id if order.stretch else None,
            passepartout_id=order.passepartout.id if order.passepartout else None,
            passepartout_length=order.passepartout_length,
            passepartout_width=order.passepartout_width,
        )
    
    # Количество копий: масштабируем расход/стоимость материалов
    copies = int(order.quantity or 1) or 1
    _scale_calc_by_copies(calculation, copies)

    # Создаем таблицу с деталями заказа
    details_table = doc.add_table(rows=1, cols=7)
    details_table.style = 'Table Grid'

    # Заголовки (компактные)
    detail_headers = ['Кол-во', 'Время', '', 'размер', 'ширина', 'цена', 'стоимость']
    detail_header_cells = details_table.rows[0].cells
    for i, header in enumerate(detail_headers):
        detail_header_cells[i].text = header
        detail_header_cells[i].paragraphs[0].runs[0].font.bold = True
        detail_header_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
        detail_header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Хелпер для служебных строк (заголовки разделов и подытоги)
    def _add_special_row(label, price=None, bold=True):
        r = details_table.add_row()
        r.cells[2].text = label
        if price is not None:
            r.cells[6].text = f"{format_number(price)} руб"
        for c in r.cells:
            for para in c.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)
                    run.font.bold = bold
        return r

    # ---------- Раздел «Материалы» ----------
    _add_special_row('МАТЕРИАЛЫ')

    # Добавляем данные о рамах
    frame_count = len(frames) if frames else 1
    has_frames = False
    for frame_idx in range(frame_count):
        if frames:
            frame_data = frames[frame_idx]
            baguette_id = frame_data.get('baguette_id')
        else:
            baguette_id = order.baguette.id if order.baguette else None
        
        if baguette_id:
            has_frames = True
            from frames.models import Baguette
            try:
                baguette = Baguette.objects.get(pk=baguette_id)
            except Baguette.DoesNotExist:
                raise ValueError(f"Багет с ID {baguette_id} не найден")
            # Размеры рамы — из frame_data при нескольких рамах, иначе из order
            frame_data_local = frames[frame_idx] if frames and frame_idx < len(frames) else {}
            rx1 = Decimal(str(frame_data_local.get('x1', order.x1))) if frame_data_local.get('x1') else order.x1
            rx2 = Decimal(str(frame_data_local.get('x2', order.x2))) if frame_data_local.get('x2') else order.x2
            if not rx1 or not rx2 or rx1 <= 0 or rx2 <= 0:
                rx1, rx2 = order.x1, order.x2
            # Ширина багета в метрах — функция ждёт метры (периметр_м + 8×ширина_м)
            baguette_quantity = PriceCalculator.calculate_baguette_quantity(
                rx1, rx2, baguette.width
            )
            # Расход и стоимость — на все копии
            baguette_quantity = baguette_quantity * copies
            baguette_price = baguette_quantity * baguette.price

            row = details_table.add_row()
            row.cells[0].text = str(copies)
            row.cells[1].text = "0.6"
            row.cells[2].text = f"Рама {frame_idx + 1}"
            row.cells[3].text = f"{format_number(rx1)}×{format_number(rx2)} см"
            # Ширина багета в модели в метрах, но в квитанции нужно показывать в см или как есть
            # Судя по примеру, ширина показывается как есть (0.048)
            row.cells[4].text = format_number(baguette.width)
            row.cells[5].text = format_number(baguette.price)
            row.cells[6].text = format_number(baguette_price)
            
            # Устанавливаем размер шрифта для всех ячеек строки
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)
            
            # Описание багета
            desc_row = details_table.add_row()
            desc_row.cells[2].text = f"Багет: {baguette.name}"
            desc_row.cells[3].text = f"расход {format_number(baguette_quantity)} м"
            for i in range(7):
                if i != 2 and i != 3:
                    desc_row.cells[i].text = ""
            
            # Устанавливаем размер шрифта для описания
            for cell in desc_row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)
        else:
            # Рама без выбранного багета — показываем по размерам (багет опционален)
            frame_data_local = frames[frame_idx] if frames and frame_idx < len(frames) else {}
            rx1 = Decimal(str(frame_data_local.get('x1', order.x1))) if frame_data_local.get('x1') else order.x1
            rx2 = Decimal(str(frame_data_local.get('x2', order.x2))) if frame_data_local.get('x2') else order.x2
            if not rx1 or not rx2 or rx1 <= 0 or rx2 <= 0:
                rx1, rx2 = order.x1, order.x2
            if rx1 and rx2 and rx1 > 0 and rx2 > 0:
                has_frames = True
                row = details_table.add_row()
                row.cells[0].text = "1"
                row.cells[1].text = "0.6"
                row.cells[2].text = f"Рама {frame_idx + 1}"
                row.cells[3].text = f"{format_number(rx1)}×{format_number(rx2)} см"
                for i in (4, 5, 6):
                    row.cells[i].text = ""
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(8)

    # Пустой заказ (нет ни рамы, ни размеров) — нечего печатать
    if not has_frames:
        raise ValueError("Заказ пустой: нет ни рамы, ни размеров")
    
    # Добавляем остальные компоненты из расчета
    components_map = {
        'glass': ('СТЕКЛО:', 'area', 'кв.м'),
        'backing': ('ПОДКЛАДКА:', None, None),
        'hardware': ('ФУРНИТУРА:', 'quantity', 'шт'),
        'podramnik': ('ПОДРАМНИК:', None, None),
        'package': ('УПАКОВКА:', 'quantity', 'шт'),
        'molding': ('МОЛДИНГ:', 'consumption', 'м'),
        'trosik': ('ТРОСИК:', 'length', 'м'),
        'podveski': ('ПОДВЕСКИ:', 'quantity', 'шт'),
        'passepartout': ('ПАСПАРТУ:', None, None),
        'stretch': ('НАТЯЖКА:', 'area', 'кв.м'),
        'work': ('РАБОТА:', None, None),
    }
    
    for component_key, (label, quantity_key, unit) in components_map.items():
        if component_key in calculation.get('components', {}):
            component = calculation['components'][component_key]
            row = details_table.add_row()
            row.cells[2].text = label
            row.cells[3].text = component.get('name', '')
            
            # Добавляем количество/расход если есть
            if quantity_key and quantity_key in component:
                quantity_value = component[quantity_key]
                row.cells[4].text = f"расход {format_number(quantity_value)} {unit}"
            elif component_key == 'hardware' and 'quantity' in component:
                row.cells[4].text = f"расход {component['quantity']} {unit}"
            
            row.cells[6].text = f"{format_number(component.get('total_price', 0))} руб"
            
            # Устанавливаем размер шрифта для всех ячеек строки
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)
    
    # Паспарту: по рамам (passepartout_frame*) и независимые (passepartout_1, ...)
    for key in sorted(calculation.get('components', {}).keys()):
        if key.startswith('passepartout_'):
            component = calculation['components'][key]
            row = details_table.add_row()
            row.cells[2].text = "ПАСПАРТУ:"
            row.cells[3].text = component.get('name', '')
            row.cells[6].text = f"{format_number(component.get('total_price', 0))} руб"
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)

    # Дополнительные подкладки (backing_2, backing_3, ...)
    for key in sorted(calculation.get('components', {}).keys()):
        if key.startswith('backing_'):
            component = calculation['components'][key]
            row = details_table.add_row()
            row.cells[2].text = "ПОДКЛАДКА:"
            row.cells[3].text = component.get('name', '')
            row.cells[6].text = f"{format_number(component.get('total_price', 0))} руб"
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)

    # Подытог по материалам (calculation['total_price'] = сумма материалов)
    materials_total = Decimal(str(calculation.get('total_price', 0)))
    _add_special_row('ИТОГО МАТЕРИАЛОВ:', price=materials_total)

    # ---------- Раздел «Работы» ----------
    _add_special_row('РАБОТЫ')

    # Работы (входят в стоимость) — по данным справочника технологических операций
    extras = OrderExtrasCalculator.for_order(order, frames)
    _scale_extras_by_copies(extras, copies)
    works_total = Decimal('0')
    for work in extras['works']['items']:
        works_total += Decimal(str(work['total']))
        row = details_table.add_row()
        row.cells[2].text = "РАБОТА:"
        row.cells[3].text = work['name']
        row.cells[6].text = f"{format_number(work['total'])} руб"
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)

    # Ручная сложность
    if extras.get('manual_complexity', 0) > 0:
        works_total += Decimal(str(extras['manual_complexity']))
        row = details_table.add_row()
        row.cells[2].text = "СЛОЖНОСТЬ:"
        row.cells[6].text = f"{format_number(extras['manual_complexity'])} руб"
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)

    # Подытог по работам
    _add_special_row('ИТОГО РАБОТ:', price=works_total)

    # Количество копий и пакетов (справочно)
    if (order.quantity or 1) > 1:
        row = details_table.add_row()
        row.cells[2].text = "КОЛИЧЕСТВО КОПИЙ:"
        row.cells[3].text = str(order.quantity)
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)
    # Количество упаковки показывается в строке «УПАКОВКА» (компонент package).

    add_table_border(details_table)
    
    # Уменьшаем отступы в таблице
    for row in details_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(1)
                para.paragraph_format.space_before = Pt(1)
    
    # Итого (компактно)
    total_para = doc.add_paragraph(f'ИТОГО ПО ЗАКАЗУ: {format_number(order.total_price)} руб')
    total_para.runs[0].font.bold = True
    total_para.runs[0].font.size = Pt(11)
    total_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_para.paragraph_format.space_before = Pt(3)
    total_para.paragraph_format.space_after = Pt(2)

    # Крупный номер заказа — приклеивается на обратную сторону рамы
    num_label = doc.add_paragraph('Заказ №')
    num_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    num_label.runs[0].font.size = Pt(11)
    num_label.paragraph_format.space_before = Pt(18)
    num_label.paragraph_format.space_after = Pt(0)
    num_para = doc.add_paragraph(str(order.pk))
    num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    num_para.runs[0].font.bold = True
    num_para.runs[0].font.size = Pt(96)
    num_para.paragraph_format.space_before = Pt(0)

    # Комментарий к заказу (если есть)
    if order.comment and order.comment.strip():
        comment_para = doc.add_paragraph()
        label_run = comment_para.add_run('Комментарий: ')
        label_run.font.bold = True
        label_run.font.size = Pt(9)
        text_run = comment_para.add_run(order.comment.strip())
        text_run.font.size = Pt(9)
        comment_para.paragraph_format.space_before = Pt(3)
        comment_para.paragraph_format.space_after = Pt(2)
    
    # Подписи заказчика/столяра и схема рам убраны по требованию.

    # Сохраняем в BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io


def generate_receipt_html(order_id):
    """
    Генерирует HTML версию квитанции для печати (та же структура, что и Word).
    Возвращает полную HTML-страницу с авто-печатью.
    """
    try:
        order = Order.objects.get(pk=order_id)
    except Order.DoesNotExist:
        raise ValueError(f"Заказ с ID {order_id} не найден")

    date_str = order.created_at.strftime('%d.%m.%Y')
    fulfillment_date_str = "—"
    if order.fulfillment_date:
        fulfillment_date_str = order.fulfillment_date.strftime('%d.%m.%Y')
    elif order.status in ['ready', 'issued']:
        fulfillment_date_str = order.updated_at.strftime('%d.%m.%Y')

    frames = []
    if order.frames_data:
        try:
            frames = json.loads(order.frames_data)
        except (json.JSONDecodeError, TypeError):
            pass

    # Расчёт стоимости (дублируем логику из Word)
    if frames and len(frames) > 0:
        calculation = {'components': {}, 'total_price': Decimal('0')}
        frame_sizes = []
        for idx, frame_data in enumerate(frames):
            if frame_data.get('baguette_id'):
                fx1 = Decimal(str(frame_data.get('x1', order.x1))) if frame_data.get('x1') else order.x1
                fx2 = Decimal(str(frame_data.get('x2', order.x2))) if frame_data.get('x2') else order.x2
                if not fx1 or not fx2 or fx1 <= 0 or fx2 <= 0:
                    fx1, fx2 = order.x1, order.x2
                frame_sizes.append((fx1, fx2))
                frame_calc = PriceCalculator.calculate_total_price(
                    x1=fx1, x2=fx2,
                    baguette_id=frame_data.get('baguette_id'),
                    passepartout_id=frame_data.get('passepartout_id'),
                    passepartout_length=Decimal(str(frame_data.get('passepartout_length'))) if frame_data.get('passepartout_length') else None,
                    passepartout_width=Decimal(str(frame_data.get('passepartout_width'))) if frame_data.get('passepartout_width') else None,
                    work_id=frame_data.get('work_id'),
                )
                frame_num = idx + 1
                for key, value in frame_calc.get('components', {}).items():
                    if key in ['baguette', 'passepartout', 'work']:
                        component_key = f'{key}_frame{frame_num}'
                        calculation['components'][component_key] = {
                            **value, 'name': f"{value.get('name', key)} (Рама {frame_num})",
                        }
                        calculation['total_price'] += Decimal(str(value.get('total_price', 0)))
        if not frame_sizes:
            frame_sizes = [(order.x1, order.x2)]
        # Подкладка/подрамник — по раме меньшего размера (на картину, один раз)
        eff_x1, eff_x2 = min(frame_sizes, key=lambda s: s[0] * s[1])
        total_glass_area = sum(PriceCalculator.calculate_glass_area(fx1, fx2) for fx1, fx2 in frame_sizes)
        other_calc = PriceCalculator.calculate_total_price(
            x1=eff_x1, x2=eff_x2,
            glass_id=None,
            backing_id=order.backing.id if order.backing else None,
            backing_ids=_order_backing_ids(order, frames),
            hardware_id=order.hardware.id if order.hardware else None,
            hardware_quantity=order.hardware_quantity or 1,
            podramnik_id=order.podramnik.id if order.podramnik else None,
            package_id=order.package.id if order.package else None,
            package_quantity=order.package_quantity or 1,
            molding_id=order.molding.id if order.molding else None,
            molding_consumption=order.molding_consumption,
            trosik_id=order.trosik.id if order.trosik else None,
            trosik_length=order.trosik_length,
            podveski_id=order.podveski.id if order.podveski else None,
            podveski_quantity=order.podveski_quantity,
        )
        for key, value in other_calc.get('components', {}).items():
            if key not in ('glass', 'stretch'):
                calculation['components'][key] = value
                calculation['total_price'] += Decimal(str(value.get('total_price', 0)))
        if order.glass and total_glass_area > 0:
            glass = order.glass
            calculation['components']['glass'] = {
                'name': glass.name,
                'area': float(total_glass_area),
                'unit_price': float(glass.price_per_sqm),
                'total_price': float(total_glass_area * glass.price_per_sqm)
            }
            calculation['total_price'] += Decimal(str(calculation['components']['glass']['total_price']))
        # Независимые паспарту (отдельный список) — входят в цену, выводим их строками
        _add_independent_passepartouts(calculation, order, frames)
        # Натяжка — работа мастера (по периметру), выводится строкой «РАБОТА» из extras,
        # как материал по площади здесь не считается.
    else:
        calculation = PriceCalculator.calculate_total_price(
            x1=order.x1, x2=order.x2,
            baguette_id=order.baguette.id if order.baguette else None,
            glass_id=order.glass.id if order.glass else None,
            backing_id=order.backing.id if order.backing else None,
            backing_ids=_order_backing_ids(order, frames),
            hardware_id=order.hardware.id if order.hardware else None,
            hardware_quantity=order.hardware_quantity or 1,
            podramnik_id=order.podramnik.id if order.podramnik else None,
            package_id=order.package.id if order.package else None,
            package_quantity=order.package_quantity or 1,
            molding_id=order.molding.id if order.molding else None,
            molding_consumption=order.molding_consumption,
            trosik_id=order.trosik.id if order.trosik else None,
            trosik_length=order.trosik_length,
            podveski_id=order.podveski.id if order.podveski else None,
            podveski_quantity=order.podveski_quantity,
            stretch_id=order.stretch.id if order.stretch else None,
            passepartout_id=order.passepartout.id if order.passepartout else None,
            passepartout_length=order.passepartout_length,
            passepartout_width=order.passepartout_width,
        )

    # Количество копий: масштабируем расход/стоимость материалов
    copies = int(order.quantity or 1) or 1
    _scale_calc_by_copies(calculation, copies)

    frame_count = len(frames) if frames else 1
    components_map = {
        'glass': ('СТЕКЛО:', 'area', 'кв.м'),
        'backing': ('ПОДКЛАДКА:', None, None),
        'hardware': ('ФУРНИТУРА:', 'quantity', 'шт'),
        'podramnik': ('ПОДРАМНИК:', None, None),
        'package': ('УПАКОВКА:', 'quantity', 'шт'),
        'molding': ('МОЛДИНГ:', 'consumption', 'м'),
        'trosik': ('ТРОСИК:', 'length', 'м'),
        'podveski': ('ПОДВЕСКИ:', 'quantity', 'шт'),
        'passepartout': ('ПАСПАРТУ:', None, None),
        'stretch': ('НАТЯЖКА:', 'area', 'кв.м'),
        'work': ('РАБОТА:', None, None),
    }

    from frames.models import Baguette
    detail_rows = []
    detail_rows.append({'section': True, 'col2': 'МАТЕРИАЛЫ'})
    has_frames = False
    for frame_idx in range(frame_count):
        if frames:
            frame_data = frames[frame_idx]
            baguette_id = frame_data.get('baguette_id')
        else:
            baguette_id = order.baguette.id if order.baguette else None
        if baguette_id:
            has_frames = True
            baguette = Baguette.objects.get(pk=baguette_id)
            frame_data_local = frames[frame_idx] if frames and frame_idx < len(frames) else {}
            rx1 = Decimal(str(frame_data_local.get('x1', order.x1))) if frame_data_local.get('x1') else order.x1
            rx2 = Decimal(str(frame_data_local.get('x2', order.x2))) if frame_data_local.get('x2') else order.x2
            if not rx1 or not rx2 or rx1 <= 0 or rx2 <= 0:
                rx1, rx2 = order.x1, order.x2
            # Ширина багета хранится в метрах — формула ждёт метры (периметр_м + 8×ширина_м)
            baguette_quantity = PriceCalculator.calculate_baguette_quantity(rx1, rx2, baguette.width) * copies
            baguette_price = baguette_quantity * baguette.price
            detail_rows.append({
                'cells': [str(copies), '0.6', f'Рама {frame_idx + 1}', f'{format_number(rx1)}×{format_number(rx2)} см',
                          format_number(baguette.width), format_number(baguette.price), format_number(baguette_price)],
                'desc': {'col2': f'Багет: {baguette.name}', 'col3': f'расход {format_number(baguette_quantity)} м'}
            })
        else:
            # Рама без выбранного багета — показываем по размерам (багет опционален)
            frame_data_local = frames[frame_idx] if frames and frame_idx < len(frames) else {}
            rx1 = Decimal(str(frame_data_local.get('x1', order.x1))) if frame_data_local.get('x1') else order.x1
            rx2 = Decimal(str(frame_data_local.get('x2', order.x2))) if frame_data_local.get('x2') else order.x2
            if not rx1 or not rx2 or rx1 <= 0 or rx2 <= 0:
                rx1, rx2 = order.x1, order.x2
            if rx1 and rx2 and rx1 > 0 and rx2 > 0:
                has_frames = True
                detail_rows.append({
                    'cells': ['1', '0.6', f'Рама {frame_idx + 1}', f'{format_number(rx1)}×{format_number(rx2)} см',
                              '', '', ''],
                    'desc': None
                })
    if not has_frames:
        raise ValueError("Заказ пустой: нет ни рамы, ни размеров")

    for component_key, (label, quantity_key, unit) in components_map.items():
        if component_key in calculation.get('components', {}):
            c = calculation['components'][component_key]
            col4 = ''
            if quantity_key and quantity_key in c:
                col4 = f"расход {format_number(c[quantity_key])} {unit}"
            elif component_key == 'hardware' and 'quantity' in c:
                col4 = f"расход {c['quantity']} {unit}"
            detail_rows.append({'component': True, 'col2': label, 'col3': c.get('name', ''), 'col4': col4, 'col6': format_number(c.get('total_price', 0))})
    for key in sorted(calculation.get('components', {}).keys()):
        if key.startswith('passepartout_'):
            c = calculation['components'][key]
            detail_rows.append({'component': True, 'col2': 'ПАСПАРТУ:', 'col3': c.get('name', ''), 'col6': format_number(c.get('total_price', 0))})
    # Дополнительные подкладки (backing_2, backing_3, ...)
    for key in sorted(calculation.get('components', {}).keys()):
        if key.startswith('backing_'):
            c = calculation['components'][key]
            detail_rows.append({'component': True, 'col2': 'ПОДКЛАДКА:', 'col3': c.get('name', ''), 'col6': format_number(c.get('total_price', 0))})
    # Подытог по материалам и переход к разделу «Работы»
    materials_total = Decimal(str(calculation.get('total_price', 0)))
    detail_rows.append({'subtotal': True, 'col2': 'ИТОГО МАТЕРИАЛОВ:', 'col6': format_number(materials_total)})
    detail_rows.append({'section': True, 'col2': 'РАБОТЫ'})

    # Работы (входят в стоимость) — по данным справочника технологических операций
    extras = OrderExtrasCalculator.for_order(order, frames)
    _scale_extras_by_copies(extras, copies)
    works_total = Decimal('0')
    for work in extras['works']['items']:
        works_total += Decimal(str(work['total']))
        detail_rows.append({'component': True, 'col2': 'РАБОТА:', 'col3': work['name'], 'col6': format_number(work['total'])})
    if extras.get('manual_complexity', 0) > 0:
        works_total += Decimal(str(extras['manual_complexity']))
        detail_rows.append({'component': True, 'col2': 'СЛОЖНОСТЬ:', 'col6': format_number(extras['manual_complexity'])})
    detail_rows.append({'subtotal': True, 'col2': 'ИТОГО РАБОТ:', 'col6': format_number(works_total)})
    if (order.quantity or 1) > 1:
        detail_rows.append({'component': True, 'col2': 'КОЛИЧЕСТВО КОПИЙ:', 'col3': str(order.quantity)})
    # Количество упаковки показывается в строке «УПАКОВКА» (компонент package).

    customer_parts = []
    if order.customer_name:
        customer_parts.append(order.customer_name)
    if order.customer_phone:
        customer_parts.append(order.customer_phone)
    customer_info = 'Клиент: ' + ' '.join(customer_parts) if customer_parts else 'Клиент: —'

    # Схема рам убрана по требованию.

    def esc(s):
        return str(s).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    html_rows = ''
    for r in detail_rows:
        if r.get('section'):
            html_rows += f'''<tr style="background:#f0f0f0;font-weight:bold"><td colspan="7">{esc(r.get('col2',''))}</td></tr>'''
        elif r.get('subtotal'):
            html_rows += f'''<tr style="font-weight:bold"><td></td><td></td><td>{esc(r.get('col2',''))}</td><td></td><td></td><td></td><td class="right">{esc(r.get('col6',''))} руб</td></tr>'''
        elif 'cells' in r:
            html_rows += f'''<tr><td>{esc(r['cells'][0])}</td><td>{esc(r['cells'][1])}</td><td>{esc(r['cells'][2])}</td>
                <td>{esc(r['cells'][3])}</td><td>{esc(r['cells'][4])}</td><td>{esc(r['cells'][5])}</td><td>{esc(r['cells'][6])}</td></tr>'''
            if r.get('desc'):
                html_rows += f'''<tr><td></td><td></td><td>{esc(r['desc']['col2'])}</td><td>{esc(r['desc']['col3'])}</td><td></td><td></td><td></td></tr>'''
        else:
            html_rows += f'''<tr><td></td><td></td><td>{esc(r.get('col2',''))}</td><td>{esc(r.get('col3',''))}</td>
                <td>{esc(r.get('col4',''))}</td><td></td><td>{esc(r.get('col6',''))} руб</td></tr>'''

    html = f'''<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>Квитанция №{order_id}</title>
<style>
body {{ font-family: Arial, sans-serif; font-size: 11pt; padding: 15px; max-width: 600px; margin: 0 auto; }}
table {{ border-collapse: collapse; width: 100%; font-size: 9pt; margin: 8px 0; }}
th, td {{ border: 1px solid #000; padding: 4px 8px; }}
th {{ text-align: center; font-weight: bold; }}
.two-col {{ display: flex; justify-content: space-between; margin: 10px 0; }}
.right {{ text-align: right; }}
.center {{ text-align: center; }}
.ip-info {{ font-size: 8pt; text-align: right; line-height: 1.3; }}
.dashed {{ border-top: 1px dashed #000; margin: 15px 0; text-align: center; font-size: 8pt; font-style: italic; }}
.total {{ font-weight: bold; font-size: 11pt; margin-top: 10px; text-align: right; }}
.signs {{ display: flex; justify-content: space-between; margin-top: 30px; }}
.sign-box {{ width: 45%; }}
@media print {{ body {{ padding: 0; }} }}
</style>
</head>
<body>
<table style="border:none"><tr><td style="border:none;vertical-align:top">
<p class="center" style="font-weight:bold;font-size:11pt">Квитанция заказа № {order.pk} от {date_str}</p>
<p><b>Сумма: {format_number(order.total_price)} руб</b></p>
<p>Аванс: {format_number(order.advance_payment)} руб</p>
<p>Долг: {format_number(order.debt)} руб</p>
</td><td style="border:none;text-align:right;vertical-align:top" class="ip-info">
ИП Караковский С.М.<br>ОГРН ИП:304263203300215<br>ИНН:263204326063<br>
Профессиональное оформление<br>картин, постеров, фотографий,<br>гобеленов и вышивок.<br>
г.Пятигорск, ул.Дзержинского, д.49А<br>тел: 33-71-75, +7-962-41-30-631<br>11:00 - 18:00. Выходной: понедельник.
</td></tr></table>
<div class="dashed">──────────────────────────────────────────────────<br>ОТРЫВНАЯ ЧАСТЬ<br>──────────────────────────────────────────────────</div>
<table style="border:none"><tr><td style="border:none;vertical-align:top">
<p class="center" style="font-weight:bold;font-size:11pt">Квитанция заказа № {order.pk} от {date_str}</p>
{f'<p class="center" style="font-size:9pt">Срок: {fulfillment_date_str}</p>' if fulfillment_date_str != "—" else ''}
</td><td style="border:none;vertical-align:top;font-size:9pt">
{customer_info} | Аванс: {format_number(order.advance_payment)} руб | <b>Сумма: {format_number(order.total_price)} руб</b>
</td></tr></table>
<p style="font-weight:bold;font-size:10pt">Пункты заказа:</p>
<table>
<tr><th>Кол-во</th><th>Время</th><th></th><th>размер</th><th>ширина</th><th>цена</th><th>стоимость</th></tr>
{html_rows}
</table>
<p class="total">ИТОГО ПО ЗАКАЗУ: {format_number(order.total_price)} руб</p>
{f'<p style="margin-top:10px;font-size:9pt"><b>Комментарий:</b> {esc(order.comment.strip())}</p>' if order.comment and order.comment.strip() else ''}
<div style="margin-top:24px;text-align:center;border-top:2px solid #000;padding-top:8px">
<div style="font-size:11pt">Заказ №</div>
<div style="font-size:96pt;font-weight:bold;line-height:1">{order.pk}</div>
</div>
<script>window.onload=function(){{window.print();}}</script>
</body>
</html>'''
    return html
